#!/usr/bin/env python
# coding: utf-8

# In[2]:


pip install pytesseract pdfplumber pillow pdf2image python-docx


# In[3]:


get_ipython().system('pip install python-docx pytesseract pdfplumber pillow pdf2image')


# In[36]:


pip install pytesseract pillow python-docx


# In[45]:


# --- célula Python: defina aqui sua função ou importe o script ---
from seu_script import imagem_para_docx

# Chama a função direto no Python
resultado = imagem_para_docx(r"C:\Users\lost4\OneDrive\Documentos\DATA\Data Transcriptions\Deloite.pdf")
print(resultado)


# In[43]:


get_ipython().system('python seu_script.py')


# In[32]:


# Path to you# Path to your scanned PDF file
pdf_path = r"C:\Users\lost4\OneDrive\Documentos\DATA\Data Transcriptions\Deloite.pdf" 


# In[26]:


import os
from pdf2image import convert_from_path

# 1) Real path to your PDF
pdf_path = r"C:\Users\lost4\OneDrive\Documentos\DATA\Data Transcriptions\Deloite.pdf"

# 2) Poppler bin folder must contain pdfinfo.exe AND pdftoppm.exe
poppler_bin = r"C:\Users\lost4\OneDrive\Documentos\DATA\Data Transcriptions\poppler-24.08.0\Library\bin"

# 3) Quick sanity checks:
if not os.path.exists(pdf_path):
    raise FileNotFoundError(f"PDF not found at {pdf_path}")

for exe in ("pdfinfo.exe","pdftoppm.exe"):
    path_exe = os.path.join(poppler_bin, exe)
    if not os.path.exists(path_exe):
        raise FileNotFoundError(f"{exe} not found in {poppler_bin}")

# 4) Convert PDF to PIL images
pages = convert_from_path(
    pdf_path,
    dpi=300,
    poppler_path=poppler_bin
)

print(f"Successfully converted PDF to {len(pages)} image(s)")


# In[51]:


import os
from PIL import Image
import pytesseract
from docx import Document

def imagem_para_docx(imagem_path):
    """
    Extrai texto de uma imagem via OCR e salva em um .docx
    Retorna mensagem de sucesso ou erro.
    """
    try:
        # 1) Verifica existência do arquivo
        if not os.path.exists(imagem_path):
            return f"Erro: arquivo não encontrado em {imagem_path}"
        
        # 2) Abre a imagem
        img = Image.open(imagem_path)
        
        # 3) Executa OCR (ajuste 'lang' se necessário: 'por', 'spa', etc.)
        texto = pytesseract.image_to_string(img, lang='eng')
        
        # 4) Prepara o documento Word
        doc = Document()
        
        if texto.strip():
            # Separa em parágrafos simples por duplo newline
            for par in texto.split('\n\n'):
                par = par.strip()
                if par:
                    doc.add_paragraph(par)
        else:
            doc.add_paragraph("(Nenhum texto detectado na imagem)")
        
        # 5) Gera nome de saída baseado no nome da imagem
        base = os.path.splitext(os.path.basename(imagem_path))[0]
        out_name = f"{base}_ocr.docx"
        doc.save(out_name)
        
        return f"Sucesso: documento gerado em {out_name}"
    
    except Exception as e:
        return f"Falha no OCR ou na geração do Word: {e}"

if __name__ == "__main__":
    # Exemplo de uso: altere para o caminho da sua imagem
    caminho_imagem = r"C:\Users\lost4\OneDrive\Documentos\DATA\Data Transcriptions\Deloite.jpg"
    resultado = imagem_para_docx(caminho_imagem)
    print(resultado)


# In[53]:


from IPython.display import FileLink
# Assuming the docx was saved as “Deloite_ocr.docx” in your current working dir
display(FileLink("Deloite_ocr.docx"))


# In[55]:


import pytesseract

# Windows example: adjust if you installed elsewhere
pytesseract.pytesseract.tesseract_cmd = (
    r"C:\Users\lost4\OneDrive\Documentos\DATA\Data Transcriptions\tesseract-ocr-w64-setup-5.5.0.20241111.exe"
)


# In[54]:


# Path to you# Path to your scanned PDF file
pdf_path = r"C:\Users\lost4\OneDrive\Documentos\DATA\Data Transcriptions\TPL PROJECT, THE BEAUTIFUL HEART .pdf" 


# In[53]:


import os
from pdf2image import convert_from_path

# 1) Ajuste aqui para o seu PDF real e para o bin do poppler
pdf_path    = r"C:\Users\lost4\OneDrive\Documentos\DATA\Data Transcriptions\TPL PROJECT, THE BEAUTIFUL HEART .pdf"
poppler_bin = r"C:\Users\lost4\OneDrive\Documentos\DATA\Data Transcriptions\poppler-24.08.0\Library\bin"

# 2) Checagens básicas
if not os.path.isfile(pdf_path):
    raise FileNotFoundError(f"PDF não encontrado em {pdf_path}")

for exe in ("pdfinfo.exe", "pdftoppm.exe"):
    p = os.path.join(poppler_bin, exe)
    if not os.path.exists(p):
        raise FileNotFoundError(f"{exe} não encontrado em {poppler_bin}")

# 3) Tenta converter para imagens
pages = convert_from_path(pdf_path, dpi=300, poppler_path=poppler_bin)
print(f"Conversão bem-sucedida: {len(pages)} página(s) detectada(s)")


# In[56]:


import os
from pdf2image import convert_from_path
import pytesseract
from docx import Document

# Aponte para o executável do Tesseract, se não estiver no PATH
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

def pdf_para_docx(pdf_path, poppler_path, lang='por'):
    # 1) Validar arquivos
    if not os.path.isfile(pdf_path):
        return f"Erro: PDF não encontrado em {pdf_path}"
    for exe in ("pdfinfo.exe","pdftoppm.exe"):
        if not os.path.exists(os.path.join(poppler_path, exe)):
            return f"Erro: {exe} não encontrado em {poppler_path}"

    # 2) Converter PDF → lista de imagens
    pages = convert_from_path(pdf_path, dpi=300, poppler_path=poppler_path)

    # 3) Criar docx e inserir OCR de cada página
    doc = Document()
    for i, img in enumerate(pages):
        if i > 0:
            doc.add_page_break()
        texto = pytesseract.image_to_string(img, lang=lang)
        if texto.strip():
            for par in texto.split('\n\n'):
                par = par.strip()
                if par:
                    doc.add_paragraph(par)
        else:
            doc.add_paragraph("(Nenhum texto detectado nesta página)")

    # 4) Salvar
    base = os.path.splitext(os.path.basename(pdf_path))[0]
    out = f"{base}_ocr.docx"
    doc.save(out)
    return f"Sucesso: documento gerado em {out}"

if __name__ == "__main__":
    pdf_path    = r"C:\Users\lost4\OneDrive\Documentos\DATA\Data Transcriptions\TPL PROJECT, THE BEAUTIFUL HEART .pdf"
    poppler_bin = r"C:\Users\lost4\OneDrive\Documentos\DATA\Data Transcriptions\poppler-24.08.0\Library\bin"

    print(pdf_para_docx(pdf_path, poppler_bin, lang='por'))


# In[ ]:


#!/usr/bin/env python3
import sys
import argparse
from pathlib import Path

import pytesseract
from pdf2image import convert_from_path
from docx import Document

# Se o tesseract.exe não estiver no PATH, descomente e ajuste:
# pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

def pdf_para_docx(pdf_path: Path, poppler_path: Path, lang: str = "por") -> Path:
    """
    Converte PDF escaneado em .docx via OCR.
    Retorna o Path para o documento gerado.
    Lança FileNotFoundError ou RuntimeError em caso de erro.
    """
    # 1) Validações
    if not pdf_path.is_file():
        raise FileNotFoundError(f"PDF não encontrado: {pdf_path}")
    if not poppler_path.is_dir():
        raise FileNotFoundError(f"Pasta Poppler não encontrada: {poppler_path}")

    for exe in ("pdfinfo.exe", "pdftoppm.exe"):
        exe_path = poppler_path / exe
        if not exe_path.exists():
            raise FileNotFoundError(f"'{exe}' não encontrado em {poppler_path}")

    # 2) Converter PDF → imagens
    try:
        pages = convert_from_path(
            str(pdf_path),
            dpi=300,
            poppler_path=str(poppler_path)
        )
    except Exception as e:
        raise RuntimeError(f"Falha ao converter PDF em imagem: {e}")

    if not pages:
        raise RuntimeError("Nenhuma página foi convertida. Verifique o PDF e o poppler_path.")

    # 3) OCR + montagem do DOCX
    doc = Document()
    for idx, img in enumerate(pages):
        if idx > 0:
            doc.add_page_break()
        texto = pytesseract.image_to_string(img, lang=lang).strip()
        if texto:
            # quebra em blocos por dupla newline
            for block in texto.split("\n\n"):
                block = block.strip()
                if block:
                    doc.add_paragraph(block)
        else:
            doc.add_paragraph("(Nenhum texto detectado nesta página)")

    # 4) Gravação do arquivo
    out_name = pdf_path.with_name(pdf_path.stem + "_ocr.docx")
    doc.save(out_name)
    return out_name

def main():
    p = argparse.ArgumentParser(
        description="Extrai texto de PDF escaneado e gera um .docx via OCR"
    )
    p.add_argument("pdf", type=Path, help="Caminho para o arquivo PDF")
    p.add_argument(
        "-p", "--poppler", type=Path, required=True,
        help="Pasta 'bin' do Poppler (contendo pdfinfo.exe e pdftoppm.exe)"
    )
    p.add_argument(
        "-l", "--lang", default="por",
        help="Idioma do Tesseract (ex: 'eng','por','spa'). Padrão: por"
    )
    args = p.parse_args()

    try:
        out_path = pdf_para_docx(args.pdf, args.poppler, lang=args.lang)
        print(f"Sucesso: documento gerado em {out_path}")
        sys.exit(0)
    except Exception as e:
        print(f"Erro: {e}", file=sys.stderr)
        sys.exit(1)

if __name__ == "__main__":
    main()


# In[ ]:


pip install pdf2image pytesseract python-docx pillow


# In[ ]:


python pdf2docx.py \
  "C:\Users\lost4\OneDrive\Documentos\DATA\Data Transcriptions\TPL PROJECT, THE BEAUTIFUL HEART.pdf" \
  -p "C:\Users\lost4\OneDrive\Documentos\DATA\Data Transcriptions\poppler-24.08.0\Library\bin" \
  -l por


# In[57]:


import os
from pdf2image import convert_from_path
import pytesseract
from docx import Document

# Se o tesseract.exe não estiver no PATH, ajuste aqui:
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

def pdf_para_docx(pdf_path, poppler_path, lang='por'):
    # 1) Validações básicas
    if not os.path.isfile(pdf_path):
        return None, f"Erro: PDF não encontrado em {pdf_path}"
    for exe in ("pdfinfo.exe","pdftoppm.exe"):
        if not os.path.exists(os.path.join(poppler_path, exe)):
            return None, f"Erro: '{exe}' não encontrado em {poppler_path}"

    # 2) Converte PDF → imagens
    pages = convert_from_path(pdf_path, dpi=300, poppler_path=poppler_path)
    if not pages:
        return None, "Erro: nenhuma página foi convertida."

    # 3) Cria o .docx e faz OCR em cada página
    doc = Document()
    for i, img in enumerate(pages):
        if i > 0:
            doc.add_page_break()
        texto = pytesseract.image_to_string(img, lang=lang)
        if texto.strip():
            for par in texto.split('\n\n'):
                par = par.strip()
                if par:
                    doc.add_paragraph(par)
        else:
            doc.add_paragraph("(Nenhum texto detectado nesta página)")

    # 4) Gera nome de saída e salva
    base = os.path.splitext(os.path.basename(pdf_path))[0].strip()
    out_name = f"{base}_ocr.docx"
    doc.save(out_name)

    return out_name, None

if __name__ == "__main__":
    pdf_path = r"C:\Users\lost4\OneDrive\Documentos\DATA\Data Transcriptions\TPL PROJECT, THE BEAUTIFUL HEART.pdf"
    poppler_bin = r"C:\Users\lost4\OneDrive\Documentos\DATA\Data Transcriptions\poppler-24.08.0\Library\bin"

    out, err = pdf_para_docx(pdf_path, poppler_bin, lang='por')
    if err:
        print(err)
    else:
        print(f"Sucesso: documento gerado em {out}")

        # Se estiver no Jupyter Notebook / Lab, cria link de download:
        try:
            from IPython.display import FileLink, display
            display(FileLink(os.path.abspath(out)))
        except ImportError:
            pass

        # Se estiver no Google Colab, dispara download automático:
        try:
            from google.colab import files
            files.download(out)
        except ImportError:
            pass

